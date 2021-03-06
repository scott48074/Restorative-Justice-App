from docx import Document
import shutil
import facesheet


class Test_facesheet:

    @classmethod
    def setup_class(self):
        self.row = [1, '2015-77128059', '10/09/2015', 'accident', 'ry4927404',
                    '18', 'psb', 'william walker', '27363 oak valley center',
                    'a15', 'los palacios', 'mn', '82156', '3/18/1999',
                    '357-219-0456', 'white', 'male', 'subject_type', 'west', 1]
        self.document = Document()

    @classmethod
    def teardown_class(self):
        """ teardown any state that was previously setup with a call to
        setup_class.
        """

    def test_parse_row(self):
        expected = {'case_number': '2015-77128059',
                    'occurred_date': '10/09/2015',
                    'incident_type': 'Accident',
                    'age': '18',
                    'name': 'William Walker',
                    'address': '27363 Oak Valley Center APT: A15 Los Palacios, MN 82156',
                    'DOB': '3/18/1999',
                    'phone': '357-219-0456',
                    'race': 'White',
                    'sex': 'Male',
                    'district': 'West'}

        results = facesheet.parse_row(self.row)

        for key in expected:
            assert expected[key] == results[key]
        self.info_dict = results

    def test_district_line(self):
        info_dict = facesheet.parse_row(self.row)
        facesheet.district_line(self.document, info_dict['district'])
        results_text = [p.text for p in self.document.paragraphs]
        expected_text = [p.text for p in Document('docx_files/district.docx').paragraphs]
        assert results_text == expected_text

    def test_approval_line(self):
        facesheet.approval_line(self.document)
        results_text = [p.text for p in self.document.paragraphs]
        expected_text = [p.text for p in Document('docx_files/approval.docx').paragraphs]
        assert results_text == expected_text

    def test_case_number_line(self):
        info_dict = facesheet.parse_row(self.row)
        facesheet.case_number_line(self.document, info_dict['case_number'])
        results_text = [p.text for p in self.document.paragraphs]
        expected_text = [p.text for p in Document('docx_files/case_number.docx').paragraphs]
        assert results_text == expected_text

    def test_name_line(self):
        info_dict = facesheet.parse_row(self.row)
        facesheet.name_line(self.document, info_dict['name'])
        results_text = [p.text for p in self.document.paragraphs]
        expected_text = [p.text for p in Document('docx_files/name.docx').paragraphs]
        assert results_text == expected_text

    def test_bio_line(self):
        info_dict = facesheet.parse_row(self.row)
        facesheet.bio_line(self.document, info_dict['sex'], info_dict['race'], info_dict['DOB'], info_dict['age'])
        results_text = [p.text for p in self.document.paragraphs]
        expected_text = [p.text for p in Document('docx_files/bio.docx').paragraphs]
        assert results_text == expected_text

    def test_charge_line(self):
        facesheet.charge_line(self.document)
        results_text = [p.text for p in self.document.paragraphs]
        expected_text = [p.text for p in Document('docx_files/charge.docx').paragraphs]
        assert results_text == expected_text

    def test_address_line(self):
        info_dict = facesheet.parse_row(self.row)
        facesheet.address_line(self.document, info_dict['address'])
        results_text = [p.text for p in self.document.paragraphs]
        expected_text = [p.text for p in Document('docx_files/address.docx').paragraphs]
        assert results_text == expected_text

    def test_phone_line(self):
        info_dict = facesheet.parse_row(self.row)
        facesheet.phone_line(self.document, info_dict['phone'])
        results_text = [p.text for p in self.document.paragraphs]
        expected_text = [p.text for p in Document('docx_files/phone.docx').paragraphs]
        assert results_text == expected_text

    def test_background_line(self):
        facesheet.background_line(self.document)
        results_text = [p.text for p in self.document.paragraphs]
        expected_text = [p.text for p in Document('docx_files/background.docx').paragraphs]
        assert results_text == expected_text

    def test_last_name_first(self):
        name = 'Scott G Frasier'
        expected = 'Frasier_Scott_G'
        assert facesheet.last_name_first(name) == expected

    def test_last_name_first_suffix(self):
        name = 'Scott G Frasier jr.'
        expected = 'Frasier_jr._Scott_G'
        assert facesheet.last_name_first(name) == expected

    def test_save_face_sheet(self):
        info_dict = facesheet.parse_row(self.row)
        document = Document('docx_files/background.docx')
        facesheet.save_facesheet(document, 'docx_files', info_dict['name'], info_dict['district'], False)
        results_text = [p.text for p in self.document.paragraphs]
        expected_text = [p.text for p in Document('docx_files/results/Walker_William/Walker_William.docx').paragraphs]
        assert results_text == expected_text
        shutil.rmtree('docx_files/results')

    def test_save_face_sheet_district_sort(self):
        info_dict = facesheet.parse_row(self.row)
        document = Document('docx_files/background.docx')
        facesheet.save_facesheet(document, 'docx_files', info_dict['name'], info_dict['district'], True)
        results_text = [p.text for p in self.document.paragraphs]
        expected_text = [p.text for p in Document('docx_files/results/West/Walker_William/Walker_William.docx').paragraphs]
        assert results_text == expected_text
        shutil.rmtree('docx_files/results')

    def test_assemble_sheet(self):
        facesheet.assemble_sheet(self.row, 'docx_files', False)
        results_text = [p.text for p in self.document.paragraphs]
        expected_text = [p.text for p in Document('docx_files/results/Walker_William/Walker_William.docx').paragraphs]
        assert results_text == expected_text
        shutil.rmtree('docx_files/results')
